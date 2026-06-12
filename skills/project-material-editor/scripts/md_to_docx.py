#!/usr/bin/env python3
"""Convert Markdown to DOCX for project/report materials.

The default path uses python-docx so common report styles are controlled inside
the generated DOCX. Pandoc is available as an opt-in path when a reference DOCX
or full Markdown coverage is more important than the built-in report styles.
"""

from __future__ import annotations

import argparse
import json
import math
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass, field, replace
from pathlib import Path
from typing import Sequence

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+['\"][^'\"]*['\"])?\)")
LIST_RE = re.compile(r"^(\s*)([-*+]|\d+[.)])\s+(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
CAPTION_NUMBER = r"\d+(?:[-.．]\d+)*"
FIGURE_CAPTION_RE = re.compile(rf"^(图|Figure|Fig\.?)\s*{CAPTION_NUMBER}([.\s、:：]|$)", re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(rf"^(表|Table)\s*{CAPTION_NUMBER}([.\s、:：]|$)", re.IGNORECASE)
DEFAULT_EAST_ASIA_FONT = "仿宋"
DEFAULT_LATIN_FONT = "Times New Roman"
HEADING_EAST_ASIA_FONT = "黑体"
MAX_IMAGE_WIDTH_INCHES = 5.8
MAX_IMAGE_HEIGHT_INCHES = 4.5


class HelpFormatter(argparse.ArgumentDefaultsHelpFormatter, argparse.RawDescriptionHelpFormatter):
    """Keep examples readable while showing defaults."""


class StyleConfigError(ValueError):
    pass


class DependencyError(RuntimeError):
    pass


@dataclass(frozen=True)
class FontConfig:
    east_asia: str = DEFAULT_EAST_ASIA_FONT
    latin: str = DEFAULT_LATIN_FONT
    heading_east_asia: str = HEADING_EAST_ASIA_FONT


@dataclass(frozen=True)
class PageConfig:
    width_cm: float = 21.0
    height_cm: float = 29.7
    top_margin_cm: float = 2.54
    bottom_margin_cm: float = 2.54
    left_margin_cm: float = 3.0
    right_margin_cm: float = 3.0


@dataclass(frozen=True)
class ImageConfig:
    max_width_inches: float = MAX_IMAGE_WIDTH_INCHES
    max_height_inches: float = MAX_IMAGE_HEIGHT_INCHES


@dataclass(frozen=True)
class ParagraphConfig:
    font_size_pt: float
    alignment: str = "left"
    bold: bool = False
    italic: bool = False
    first_line_indent_cm: float | None = None
    left_indent_cm: float | None = None
    line_spacing: float | None = 1.0
    space_before_pt: float = 0.0
    space_after_pt: float = 0.0


@dataclass(frozen=True)
class HeadingConfig:
    font_size_pt: float
    alignment: str
    space_before_pt: float
    space_after_pt: float


@dataclass(frozen=True)
class TableConfig:
    font_size_pt: float = 10.0


@dataclass(frozen=True)
class StyleConfig:
    fonts: FontConfig = field(default_factory=FontConfig)
    page: PageConfig = field(default_factory=PageConfig)
    image: ImageConfig = field(default_factory=ImageConfig)
    body: ParagraphConfig = field(
        default_factory=lambda: ParagraphConfig(
            font_size_pt=11.0,
            first_line_indent_cm=0.74,
            line_spacing=1.5,
            space_after_pt=6.0,
        )
    )
    figure_image: ParagraphConfig = field(
        default_factory=lambda: ParagraphConfig(
            font_size_pt=11.0,
            alignment="center",
            line_spacing=1.0,
            space_before_pt=6.0,
            space_after_pt=3.0,
        )
    )
    missing_image: ParagraphConfig = field(
        default_factory=lambda: ParagraphConfig(
            font_size_pt=10.0,
            alignment="center",
            italic=True,
            line_spacing=1.0,
            space_before_pt=6.0,
            space_after_pt=3.0,
        )
    )
    figure_caption: ParagraphConfig = field(
        default_factory=lambda: ParagraphConfig(
            font_size_pt=10.0,
            alignment="center",
            line_spacing=1.0,
            space_after_pt=6.0,
        )
    )
    table_caption: ParagraphConfig = field(
        default_factory=lambda: ParagraphConfig(
            font_size_pt=10.0,
            alignment="center",
            bold=True,
            line_spacing=1.0,
            space_before_pt=6.0,
            space_after_pt=3.0,
        )
    )
    table_header: ParagraphConfig = field(
        default_factory=lambda: ParagraphConfig(font_size_pt=10.0, alignment="center", bold=True)
    )
    table_body: ParagraphConfig = field(default_factory=lambda: ParagraphConfig(font_size_pt=10.0))
    table_spacer: ParagraphConfig = field(default_factory=lambda: ParagraphConfig(font_size_pt=1.0, space_after_pt=6.0))
    list_item: ParagraphConfig = field(
        default_factory=lambda: ParagraphConfig(
            font_size_pt=11.0,
            first_line_indent_cm=-0.37,
            left_indent_cm=0.74,
            line_spacing=1.5,
            space_after_pt=3.0,
        )
    )
    table: TableConfig = field(default_factory=TableConfig)
    headings: tuple[HeadingConfig, ...] = field(
        default_factory=lambda: (
            HeadingConfig(16.0, "center", 12.0, 12.0),
            HeadingConfig(14.0, "left", 10.0, 6.0),
            HeadingConfig(12.0, "left", 8.0, 4.0),
            HeadingConfig(11.0, "left", 6.0, 3.0),
            HeadingConfig(11.0, "left", 6.0, 3.0),
            HeadingConfig(11.0, "left", 6.0, 3.0),
        )
    )


@dataclass(frozen=True)
class Block:
    kind: str
    text: str = ""
    level: int = 0
    target: str = ""
    rows: tuple[tuple[str, ...], ...] = ()


def run_pandoc(markdown: Path, output: Path, reference_doc: Path | None, pandoc_bin: str) -> int:
    command = [pandoc_bin, str(markdown), "-o", str(output), "--resource-path", str(markdown.parent)]
    if reference_doc:
        command.extend(["--reference-doc", str(reference_doc)])
    subprocess.run(command, check=True)
    return 0


def ensure_python_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise DependencyError(
            "python-docx is required for fallback conversion. "
            "Use a temporary run such as: uv run --with python-docx python "
            "skills/project-material-editor/scripts/md_to_docx.py input.md"
        ) from exc


def load_style_config(path: Path | None) -> StyleConfig:
    config = StyleConfig()
    if path is None:
        return config
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except OSError as exc:
        raise StyleConfigError(f"Could not read style config {path}: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise StyleConfigError(f"Invalid JSON in style config {path}: {exc}") from exc
    if not isinstance(raw, dict):
        raise StyleConfigError("Style config must be a JSON object.")

    allowed = {"fonts", "page", "image", "body", "headings", "captions", "tables", "lists"}
    reject_unknown_keys(raw, allowed, "style config")

    for key, value in raw.items():
        if key == "fonts":
            config = replace(config, fonts=update_fonts(config.fonts, value, key))
        elif key == "page":
            config = replace(config, page=update_page(config.page, value, key))
        elif key == "image":
            config = replace(config, image=update_image(config.image, value, key))
        elif key == "body":
            config = replace(config, body=update_paragraph(config.body, value, key))
        elif key == "headings":
            config = replace(config, headings=update_headings(config.headings, value, key))
        elif key == "captions":
            figure, table = update_captions(config.figure_caption, config.table_caption, value, key)
            config = replace(config, figure_caption=figure, table_caption=table)
        elif key == "tables":
            table_config, header, body = update_tables(config.table, config.table_header, config.table_body, value, key)
            config = replace(config, table=table_config, table_header=header, table_body=body)
        elif key == "lists":
            config = replace(config, list_item=update_paragraph(config.list_item, value, key))
    return config


def reject_unknown_keys(data: dict[str, object], allowed: set[str], path: str) -> None:
    unknown = sorted(set(data) - allowed)
    if unknown:
        raise StyleConfigError(f"Unknown key(s) in {path}: {', '.join(unknown)}")


def require_object(value: object, path: str) -> dict[str, object]:
    if not isinstance(value, dict):
        raise StyleConfigError(f"{path} must be an object.")
    return value


def string_value(value: object, path: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise StyleConfigError(f"{path} must be a non-empty string.")
    return value


def bool_value(value: object, path: str) -> bool:
    if not isinstance(value, bool):
        raise StyleConfigError(f"{path} must be true or false.")
    return value


def numeric_value(value: object, path: str, *, positive: bool = False, non_negative: bool = False) -> float:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        raise StyleConfigError(f"{path} must be a number.")
    result = float(value)
    if not math.isfinite(result):
        raise StyleConfigError(f"{path} must be a finite number.")
    if positive and result <= 0:
        raise StyleConfigError(f"{path} must be greater than 0.")
    if non_negative and result < 0:
        raise StyleConfigError(f"{path} must be 0 or greater.")
    return result


def alignment_value(value: object, path: str) -> str:
    result = string_value(value, path).lower()
    if result not in {"left", "center", "right", "justify"}:
        raise StyleConfigError(f"{path} must be one of: left, center, right, justify.")
    return result


def update_fonts(current: FontConfig, value: object, path: str) -> FontConfig:
    data = require_object(value, path)
    reject_unknown_keys(data, {"east_asia", "latin", "heading_east_asia"}, path)
    updates: dict[str, str] = {}
    for key, raw in data.items():
        updates[key] = string_value(raw, f"{path}.{key}")
    return replace(current, **updates)


def update_page(current: PageConfig, value: object, path: str) -> PageConfig:
    data = require_object(value, path)
    allowed = {
        "width_cm",
        "height_cm",
        "margin_cm",
        "top_margin_cm",
        "bottom_margin_cm",
        "left_margin_cm",
        "right_margin_cm",
    }
    reject_unknown_keys(data, allowed, path)
    updates: dict[str, float] = {}
    if "margin_cm" in data:
        margin = numeric_value(data["margin_cm"], f"{path}.margin_cm", positive=True)
        updates.update(
            {
                "top_margin_cm": margin,
                "bottom_margin_cm": margin,
                "left_margin_cm": margin,
                "right_margin_cm": margin,
            }
        )
    for key in allowed - {"margin_cm"}:
        if key in data:
            updates[key] = numeric_value(data[key], f"{path}.{key}", positive=True)
    return replace(current, **updates)


def update_image(current: ImageConfig, value: object, path: str) -> ImageConfig:
    data = require_object(value, path)
    reject_unknown_keys(data, {"max_width_inches", "max_height_inches"}, path)
    updates = {
        key: numeric_value(raw, f"{path}.{key}", positive=True)
        for key, raw in data.items()
    }
    return replace(current, **updates)


def update_paragraph(current: ParagraphConfig, value: object, path: str) -> ParagraphConfig:
    data = require_object(value, path)
    allowed = {
        "font_size_pt",
        "alignment",
        "bold",
        "italic",
        "first_line_indent_cm",
        "left_indent_cm",
        "line_spacing",
        "space_before_pt",
        "space_after_pt",
    }
    reject_unknown_keys(data, allowed, path)
    updates: dict[str, object] = {}
    for key, raw in data.items():
        item_path = f"{path}.{key}"
        if key == "font_size_pt":
            updates[key] = numeric_value(raw, item_path, positive=True)
        elif key == "alignment":
            updates[key] = alignment_value(raw, item_path)
        elif key in {"bold", "italic"}:
            updates[key] = bool_value(raw, item_path)
        elif key == "line_spacing":
            updates[key] = None if raw is None else numeric_value(raw, item_path, positive=True)
        elif key in {"space_before_pt", "space_after_pt"}:
            updates[key] = numeric_value(raw, item_path, non_negative=True)
        else:
            updates[key] = None if raw is None else numeric_value(raw, item_path)
    return replace(current, **updates)


def update_headings(current: tuple[HeadingConfig, ...], value: object, path: str) -> tuple[HeadingConfig, ...]:
    data = require_object(value, path)
    heading_list = list(current)
    for key, raw in data.items():
        if key not in {"1", "2", "3", "4", "5", "6"}:
            raise StyleConfigError(f"{path} keys must be heading levels 1-6.")
        heading_list[int(key) - 1] = update_heading(heading_list[int(key) - 1], raw, f"{path}.{key}")
    return tuple(heading_list)


def update_heading(current: HeadingConfig, value: object, path: str) -> HeadingConfig:
    data = require_object(value, path)
    allowed = {"font_size_pt", "alignment", "space_before_pt", "space_after_pt"}
    reject_unknown_keys(data, allowed, path)
    updates: dict[str, object] = {}
    for key, raw in data.items():
        item_path = f"{path}.{key}"
        if key == "alignment":
            updates[key] = alignment_value(raw, item_path)
        elif key == "font_size_pt":
            updates[key] = numeric_value(raw, item_path, positive=True)
        else:
            updates[key] = numeric_value(raw, item_path, non_negative=True)
    return replace(current, **updates)


def update_captions(
    figure: ParagraphConfig,
    table: ParagraphConfig,
    value: object,
    path: str,
) -> tuple[ParagraphConfig, ParagraphConfig]:
    data = require_object(value, path)
    reject_unknown_keys(data, {"figure", "table"}, path)
    if "figure" in data:
        figure = update_paragraph(figure, data["figure"], f"{path}.figure")
    if "table" in data:
        table = update_paragraph(table, data["table"], f"{path}.table")
    return figure, table


def update_tables(
    table: TableConfig,
    header: ParagraphConfig,
    body: ParagraphConfig,
    value: object,
    path: str,
) -> tuple[TableConfig, ParagraphConfig, ParagraphConfig]:
    data = require_object(value, path)
    reject_unknown_keys(data, {"font_size_pt", "header", "body"}, path)
    if "font_size_pt" in data:
        table = replace(table, font_size_pt=numeric_value(data["font_size_pt"], f"{path}.font_size_pt", positive=True))
    if "header" in data:
        header = update_paragraph(header, data["header"], f"{path}.header")
    if "body" in data:
        body = update_paragraph(body, data["body"], f"{path}.body")
    return table, header, body


def apply_cli_style_overrides(config: StyleConfig, args: argparse.Namespace) -> StyleConfig:
    fonts = config.fonts
    if args.east_asia_font:
        fonts = replace(fonts, east_asia=args.east_asia_font)
    if args.latin_font:
        fonts = replace(fonts, latin=args.latin_font)
    if args.heading_east_asia_font:
        fonts = replace(fonts, heading_east_asia=args.heading_east_asia_font)

    page = config.page
    if args.margin_cm is not None:
        page = replace(
            page,
            top_margin_cm=args.margin_cm,
            bottom_margin_cm=args.margin_cm,
            left_margin_cm=args.margin_cm,
            right_margin_cm=args.margin_cm,
        )
    for attr in ("top_margin_cm", "bottom_margin_cm", "left_margin_cm", "right_margin_cm"):
        value = getattr(args, attr)
        if value is not None:
            page = replace(page, **{attr: value})

    image = config.image
    if args.max_image_width_inches is not None:
        image = replace(image, max_width_inches=args.max_image_width_inches)
    if args.max_image_height_inches is not None:
        image = replace(image, max_height_inches=args.max_image_height_inches)

    return replace(config, fonts=fonts, page=page, image=image)


def style_config_from_args(args: argparse.Namespace) -> StyleConfig:
    return apply_cli_style_overrides(load_style_config(args.style_config), args)


def split_table_row(line: str) -> tuple[str, ...]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return tuple(cell.strip() for cell in stripped.split("|"))


def normalize_table_row(values: tuple[str, ...], width: int) -> tuple[str, ...]:
    if len(values) < width:
        return values + ("",) * (width - len(values))
    if len(values) > width:
        return values[: width - 1] + (" | ".join(values[width - 1 :]),)
    return values


def warn_uneven_table_row(row_index: int, actual: int, expected: int) -> None:
    if actual != expected:
        print(
            f"Warning: table row {row_index} has {actual} cell(s); expected {expected}.",
            file=sys.stderr,
        )


def flush_paragraph(blocks: list[Block], lines: list[str]) -> None:
    if not lines:
        return
    text = " ".join(line.strip() for line in lines if line.strip())
    if text:
        blocks.append(Block("paragraph", text=text))
    lines.clear()


def parse_blocks(markdown: Path) -> list[Block]:
    lines = markdown.read_text(encoding="utf-8").splitlines()
    blocks: list[Block] = []
    paragraph_lines: list[str] = []
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            flush_paragraph(blocks, paragraph_lines)
            index += 1
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            flush_paragraph(blocks, paragraph_lines)
            blocks.append(Block("heading", text=heading.group(2).strip(), level=min(len(heading.group(1)), 6)))
            index += 1
            continue

        image = IMAGE_RE.fullmatch(stripped)
        if image:
            flush_paragraph(blocks, paragraph_lines)
            blocks.append(Block("image", text=image.group(1).strip(), target=image.group(2).strip()))
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[index + 1]):
            flush_paragraph(blocks, paragraph_lines)
            rows = [split_table_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            blocks.append(Block("table", rows=tuple(rows)))
            continue

        list_item = LIST_RE.match(line)
        if list_item:
            flush_paragraph(blocks, paragraph_lines)
            marker = list_item.group(2)
            style = "ordered_list" if marker[0].isdigit() else "unordered_list"
            blocks.append(Block(style, text=list_item.group(3).strip()))
            index += 1
            continue

        paragraph_lines.append(stripped)
        index += 1

    flush_paragraph(blocks, paragraph_lines)
    return classify_caption_blocks(blocks)


def classify_caption_blocks(blocks: list[Block]) -> list[Block]:
    result: list[Block] = []
    for index, block in enumerate(blocks):
        if block.kind == "paragraph":
            if FIGURE_CAPTION_RE.match(block.text):
                result.append(Block("figure_caption", text=block.text))
                continue
            if TABLE_CAPTION_RE.match(block.text):
                result.append(Block("table_caption", text=block.text))
                continue
            if index + 1 < len(blocks) and blocks[index + 1].kind == "table":
                prefix, caption = split_trailing_table_caption(block.text)
                if caption:
                    if prefix:
                        result.append(Block("paragraph", text=prefix))
                    result.append(Block("table_caption", text=caption))
                    continue
        result.append(block)
    return result


def split_trailing_table_caption(text: str) -> tuple[str, str]:
    if TABLE_CAPTION_RE.match(text):
        return "", text

    markers = [
        match.start()
        for match in re.finditer(rf"(?:表|Table)\s*{CAPTION_NUMBER}([.\s、:：]|$)", text, re.IGNORECASE)
    ]
    for start in reversed(markers):
        prefix = text[:start].rstrip()
        caption = text[start:].strip()
        if not prefix:
            return "", caption
        if prefix[-1] in "。.!?！？；;":
            return prefix, caption
    return text, ""


def set_cell_text(cell: object, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.style = "Table Body"


def add_caption(document: object, text: str, caption_kind: str) -> None:
    style_name = "Table Caption" if caption_kind == "table" else "Figure Caption"
    document.add_paragraph(text, style=style_name)


def add_picture(document: object, markdown: Path, target: str, config: StyleConfig) -> bool:
    image_path = Path(target)
    if not image_path.is_absolute():
        image_path = markdown.parent / image_path
    if not image_path.exists():
        document.add_paragraph(f"[Missing image: {target}]", style="Missing Image")
        return False
    paragraph = document.add_paragraph(style="Figure Image")
    run = paragraph.add_run()
    width, height = scaled_image_size(image_path, config)
    run.add_picture(str(image_path), width=width, height=height)
    return True


def add_table(document: object, rows: tuple[tuple[str, ...], ...]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    normalized = [normalize_table_row(row, width) for row in rows]
    table = document.add_table(rows=1, cols=width)
    table.style = "Report Table"

    for cell, text in zip(table.rows[0].cells, normalized[0]):
        set_cell_text(cell, text)
        for paragraph in cell.paragraphs:
            paragraph.style = "Table Header"

    for row_index, row_values in enumerate(rows, start=1):
        warn_uneven_table_row(row_index, len(row_values), width)

    for row_values in normalized[1:]:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            set_cell_text(cell, text)

    document.add_paragraph(style="Table Spacer")


def setup_document_styles(document: object, config: StyleConfig) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_SECTION_START
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(config.page.width_cm)
    section.page_height = Cm(config.page.height_cm)
    section.top_margin = Cm(config.page.top_margin_cm)
    section.bottom_margin = Cm(config.page.bottom_margin_cm)
    section.left_margin = Cm(config.page.left_margin_cm)
    section.right_margin = Cm(config.page.right_margin_cm)

    normal = document.styles["Normal"]
    set_style_fonts(normal, config.fonts.latin, config.fonts.east_asia)
    normal.font.size = Pt(config.body.font_size_pt)
    normal.paragraph_format.first_line_indent = Cm(config.body.first_line_indent_cm or 0)
    normal.paragraph_format.line_spacing = config.body.line_spacing
    normal.paragraph_format.space_before = Pt(config.body.space_before_pt)
    normal.paragraph_format.space_after = Pt(config.body.space_after_pt)

    configure_heading_styles(document, config)
    configure_body_styles(document, WD_STYLE_TYPE, config)
    ensure_table_style(document, "Report Table", WD_STYLE_TYPE, config)


def ensure_paragraph_style(
    document: object,
    name: str,
    style_type: object,
    *,
    alignment: int,
    latin_font: str,
    east_asia_font: str,
    font_size: float,
    bold: bool,
    italic: bool,
    line_spacing: float | None = None,
    space_before: float = 0,
    space_after: float = 0,
    first_line_indent_cm: float | None = None,
    left_indent_cm: float | None = None,
    keep_next: bool = False,
    keep_together: bool = False,
) -> None:
    from docx.shared import Cm, Pt

    if name in document.styles:
        style = document.styles[name]
    else:
        style = document.styles.add_style(name, style_type)
    set_style_fonts(style, latin_font, east_asia_font)
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.font.italic = italic
    style.paragraph_format.alignment = alignment
    style.paragraph_format.first_line_indent = Cm(first_line_indent_cm) if first_line_indent_cm is not None else None
    style.paragraph_format.left_indent = Cm(left_indent_cm) if left_indent_cm is not None else None
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    if keep_next:
        set_style_keep_next(style)
    if keep_together:
        set_style_keep_together(style)


def configure_body_styles(document: object, style_type: object, config: StyleConfig) -> None:
    ensure_paragraph_style(
        document,
        "Body Text",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.body, config, alignment_lookup()),
    )
    ensure_paragraph_style(
        document,
        "Figure Image",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.figure_image, config, alignment_lookup()),
        keep_next=True,
    )
    ensure_paragraph_style(
        document,
        "Missing Image",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.missing_image, config, alignment_lookup()),
        keep_next=True,
    )
    ensure_paragraph_style(
        document,
        "Figure Caption",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.figure_caption, config, alignment_lookup()),
        keep_together=True,
    )
    ensure_paragraph_style(
        document,
        "Table Caption",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.table_caption, config, alignment_lookup()),
        keep_together=True,
    )
    ensure_paragraph_style(
        document,
        "Table Header",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.table_header, config, alignment_lookup()),
    )
    ensure_paragraph_style(
        document,
        "Table Body",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.table_body, config, alignment_lookup()),
    )
    ensure_paragraph_style(
        document,
        "Table Spacer",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.table_spacer, config, alignment_lookup()),
    )
    ensure_paragraph_style(
        document,
        "List Bullet",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.list_item, config, alignment_lookup()),
    )
    ensure_paragraph_style(
        document,
        "List Number",
        style_type.PARAGRAPH,
        **paragraph_style_kwargs(config.list_item, config, alignment_lookup()),
    )


def paragraph_style_kwargs(paragraph: ParagraphConfig, config: StyleConfig, align: dict[str, int]) -> dict[str, object]:
    return {
        "alignment": align[paragraph.alignment],
        "latin_font": config.fonts.latin,
        "east_asia_font": config.fonts.east_asia,
        "font_size": paragraph.font_size_pt,
        "bold": paragraph.bold,
        "italic": paragraph.italic,
        "line_spacing": paragraph.line_spacing,
        "space_before": paragraph.space_before_pt,
        "space_after": paragraph.space_after_pt,
        "first_line_indent_cm": paragraph.first_line_indent_cm,
        "left_indent_cm": paragraph.left_indent_cm,
    }


def alignment_lookup() -> dict[str, int]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }


def ensure_table_style(document: object, name: str, style_type: object, config: StyleConfig) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    if name in document.styles:
        style = document.styles[name]
    else:
        style = document.styles.add_style(name, style_type.TABLE)
    if "Table Grid" in document.styles:
        style.base_style = document.styles["Table Grid"]
    set_style_fonts(style, config.fonts.latin, config.fonts.east_asia)
    style.font.size = Pt(config.table.font_size_pt)

    existing = style.element.find(qn("w:tblPr"))
    if existing is not None:
        style.element.remove(existing)
    tbl_pr = OxmlElement("w:tblPr")
    tbl_pr.append(make_style_value_element("w:jc", "center"))
    tbl_pr.append(make_table_borders())
    tbl_pr.append(make_table_cell_margins())
    style.element.append(tbl_pr)


def make_style_value_element(tag: str, value: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = OxmlElement(tag)
    element.set(qn("w:val"), value)
    return element


def make_table_borders():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    return borders


def make_table_cell_margins():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    margins = OxmlElement("w:tblCellMar")
    for side, width in (("top", "80"), ("left", "108"), ("bottom", "80"), ("right", "108")):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), width)
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    return margins


def configure_heading_styles(document: object, config: StyleConfig) -> None:
    align = alignment_lookup()

    for level, heading in enumerate(config.headings, start=1):
        style_name = f"Heading {level}"
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        set_style_fonts(style, config.fonts.latin, config.fonts.heading_east_asia)
        style.font.size = pt(heading.font_size_pt)
        style.font.bold = True
        style.paragraph_format.alignment = align[heading.alignment]
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = pt(heading.space_before_pt)
        style.paragraph_format.space_after = pt(heading.space_after_pt)


def set_style_fonts(style: object, latin_font: str, east_asia_font: str) -> None:
    from docx.oxml.ns import qn

    style.font.name = latin_font
    r_fonts = style.element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_style_keep_next(style: object) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = style.element.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_style_keep_together(style: object) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = style.element.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def pt(value: float):
    from docx.shared import Pt

    return Pt(value)


def scaled_image_size(image_path: Path, config: StyleConfig):
    from docx.shared import Inches

    max_width = Inches(config.image.max_width_inches)
    max_height = Inches(config.image.max_height_inches)
    try:
        from PIL import Image

        with Image.open(image_path) as image:
            width_px, height_px = image.size
        if width_px <= 0 or height_px <= 0:
            return max_width, None
        image_width = Inches(width_px / 96)
        image_height = Inches(height_px / 96)
        scale = min(max_width / image_width, max_height / image_height, 1.0)
        return int(image_width * scale), int(image_height * scale)
    except Exception:
        return max_width, None


def render_blocks(markdown: Path, output: Path, blocks: list[Block], config: StyleConfig, strict_assets: bool) -> int:
    ensure_python_docx()
    from docx import Document

    document = Document()
    setup_document_styles(document, config)
    missing_images = 0
    index = 0

    while index < len(blocks):
        block = blocks[index]

        if block.kind == "heading":
            document.add_heading(block.text, level=block.level)
        elif block.kind == "paragraph":
            document.add_paragraph(block.text, style="Body Text")
        elif block.kind == "ordered_list":
            document.add_paragraph(block.text, style="List Number")
        elif block.kind == "unordered_list":
            document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "table_caption":
            add_caption(document, block.text, "table")
        elif block.kind == "figure_caption":
            add_caption(document, block.text, "figure")
        elif block.kind == "table":
            add_table(document, block.rows)
        elif block.kind == "image":
            if not add_picture(document, markdown, block.target, config):
                missing_images += 1
            if index + 1 < len(blocks) and blocks[index + 1].kind == "figure_caption":
                index += 1
                add_caption(document, blocks[index].text, "figure")
            elif block.text:
                add_caption(document, block.text, "figure")
        else:
            document.add_paragraph(block.text, style="Body Text")
        index += 1

    document.save(output)
    if missing_images:
        print(f"Warning: {missing_images} image(s) were missing.", file=sys.stderr)
    return 1 if strict_assets and missing_images else 0


def run_fallback(markdown: Path, output: Path, config: StyleConfig, strict_assets: bool) -> int:
    return render_blocks(markdown, output, parse_blocks(markdown), config, strict_assets)


def positive_float(value: str) -> float:
    try:
        result = float(value)
    except ValueError as exc:
        raise argparse.ArgumentTypeError("must be a number") from exc
    if not math.isfinite(result):
        raise argparse.ArgumentTypeError("must be a finite number")
    if result <= 0:
        raise argparse.ArgumentTypeError("must be greater than 0")
    return result


def build_parser() -> argparse.ArgumentParser:
    epilog = """\
Examples:
  Built-in renderer, safe by default:
    python skills/project-material-editor/scripts/md_to_docx.py docs/report.md

  Replace an existing generated DOCX intentionally:
    python skills/project-material-editor/scripts/md_to_docx.py docs/report.md --overwrite

  Use a reference DOCX template through Pandoc:
    python skills/project-material-editor/scripts/md_to_docx.py docs/report.md --reference-doc templates/reference.docx

  Tune common built-in renderer styles:
    python skills/project-material-editor/scripts/md_to_docx.py docs/report.md --east-asia-font 宋体 --margin-cm 2.5

  Use a JSON style config for built-in renderer details:
    python skills/project-material-editor/scripts/md_to_docx.py docs/report.md --style-config styles/report-docx.json

Defaults and behavior:
  - Run from the project root with project-relative paths when possible.
  - Output defaults to <markdown stem>.docx and is not overwritten unless --overwrite is passed.
  - Local images and Pandoc resource paths resolve relative to the Markdown file's directory.
  - The built-in renderer supports common report blocks: headings, paragraphs, simple lists,
    local images, figure/table captions, and pipe tables.
  - Built-in rendering requires python-docx. Pillow is optional and only improves image scaling.
  - --reference-doc selects Pandoc. If Pandoc is requested but unavailable, the command fails
    instead of silently ignoring the template.
  - Style flags and --style-config apply only to the built-in renderer; use --reference-doc
    for Pandoc styling.
  - Exit codes: 0 success, 1 conversion completed with strict asset failures or subprocess
    failure, 2 invalid input, missing dependency, unsafe overwrite, or invalid config.

Style config JSON:
  Top-level keys: fonts, page, image, body, headings, captions, tables, lists.
  CLI style flags override JSON values. Internal Word style names and OXML table details are
  intentionally not part of the config contract.
"""
    parser = argparse.ArgumentParser(
        description=__doc__,
        epilog=epilog,
        formatter_class=HelpFormatter,
    )
    parser.add_argument("markdown", type=Path, help="Markdown source file.")
    parser.add_argument("-o", "--output", type=Path, help="DOCX output path. Defaults to the Markdown stem.")
    parser.add_argument("--overwrite", action="store_true", help="Allow replacing an existing DOCX output file.")
    parser.add_argument("--reference-doc", type=Path, help="Pandoc reference DOCX. Implies --pandoc.")
    renderer = parser.add_mutually_exclusive_group()
    renderer.add_argument(
        "--pandoc",
        action="store_true",
        help="Use Pandoc instead of the built-in python-docx report renderer.",
    )
    renderer.add_argument(
        "--fallback",
        action="store_true",
        help="Use the built-in python-docx report renderer. This is the default without --pandoc or --reference-doc.",
    )
    parser.add_argument("--pandoc-bin", default="pandoc", help="Pandoc executable name or path.")
    parser.add_argument("--style-config", type=Path, help="JSON style config for the built-in renderer.")
    parser.add_argument("--east-asia-font", help="Built-in renderer East Asian body font.")
    parser.add_argument("--latin-font", help="Built-in renderer Latin body font.")
    parser.add_argument("--heading-east-asia-font", help="Built-in renderer East Asian heading font.")
    parser.add_argument("--margin-cm", type=positive_float, help="Set all page margins in centimeters.")
    parser.add_argument("--top-margin-cm", type=positive_float, help="Set top page margin in centimeters.")
    parser.add_argument("--bottom-margin-cm", type=positive_float, help="Set bottom page margin in centimeters.")
    parser.add_argument("--left-margin-cm", type=positive_float, help="Set left page margin in centimeters.")
    parser.add_argument("--right-margin-cm", type=positive_float, help="Set right page margin in centimeters.")
    parser.add_argument("--max-image-width-inches", type=positive_float, help="Built-in renderer maximum image width.")
    parser.add_argument("--max-image-height-inches", type=positive_float, help="Built-in renderer maximum image height.")
    parser.add_argument(
        "--strict-assets",
        action="store_true",
        help="Return exit code 1 when referenced local images are missing.",
    )
    return parser


def has_builtin_style_overrides(args: argparse.Namespace) -> bool:
    return any(
        getattr(args, name) is not None
        for name in (
            "style_config",
            "east_asia_font",
            "latin_font",
            "heading_east_asia_font",
            "margin_cm",
            "top_margin_cm",
            "bottom_margin_cm",
            "left_margin_cm",
            "right_margin_cm",
            "max_image_width_inches",
            "max_image_height_inches",
        )
    )


def main(argv: Sequence[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    markdown = args.markdown.resolve()
    if not markdown.exists():
        print(f"File not found: {markdown}", file=sys.stderr)
        return 2
    output = (args.output or markdown.with_suffix(".docx")).resolve()
    if output.exists() and not args.overwrite:
        print(f"Output already exists, use --overwrite to replace it: {output}", file=sys.stderr)
        return 2

    reference_doc = args.reference_doc.resolve() if args.reference_doc else None
    if reference_doc and not reference_doc.exists():
        print(f"Reference DOCX not found: {reference_doc}", file=sys.stderr)
        return 2
    if args.fallback and reference_doc:
        print("--fallback cannot be used with --reference-doc because reference DOCX requires Pandoc.", file=sys.stderr)
        return 2

    output.parent.mkdir(parents=True, exist_ok=True)
    use_pandoc = args.pandoc or reference_doc is not None
    if use_pandoc:
        if has_builtin_style_overrides(args):
            print("Warning: built-in style options are ignored when using Pandoc.", file=sys.stderr)
        pandoc_path = shutil.which(args.pandoc_bin)
        if not pandoc_path:
            print(
                f"Pandoc was requested but '{args.pandoc_bin}' is not available; install Pandoc or use --fallback.",
                file=sys.stderr,
            )
            return 2
        try:
            return run_pandoc(markdown, output, reference_doc, pandoc_path)
        except subprocess.CalledProcessError as exc:
            print(f"Pandoc failed with exit code {exc.returncode}.", file=sys.stderr)
            return 1

    try:
        config = style_config_from_args(args)
        return run_fallback(markdown, output, config, args.strict_assets)
    except (DependencyError, StyleConfigError) as exc:
        print(str(exc), file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
